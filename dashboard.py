import os
import re
import sqlite3
import requests
import datetime
import io
import hashlib
import html
import json
import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import seaborn as sns

# ==============================================================================
# 1. PAGE CONFIGURATION & PREMIUM HIGH-CONTRAST UI INJECTION
# ==============================================================================
st.set_page_config(
    page_title="Agentic AI for Smart Facility Operations and Optimization Dashboard",
    page_icon="⚙️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Resolve paths relative to this script's location so the app doesn't break
# depending on which directory it's launched from.
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_FILE = os.path.join(BASE_DIR, "ai4i2020.csv")
DB_FILE = os.path.join(BASE_DIR, "work_orders.db")

# COLOR PALETTE:
# App Background: #0B1121 (Deep Space Blue)
# Sidebar/Forms: #0F1423
# Cards: #151B2E
# Primary Accent: #00D2FF (Electric Cyan)
# Secondary Accent: #3A86FF (Bright Blue)
# Success Accent: #00FF9D (Spring Green)
# Warning/Alert: #FF007A (Neon Pink)
# Text Primary: #FFFFFF
# Text Muted: #9CA3AF

def inject_enterprise_ui():
    """Overrides default Streamlit styling with a sleek, high-contrast Deep Space UI."""
    st.markdown("""
    <style>
        /* Modern Scrollbar */
        ::-webkit-scrollbar { width: 8px; height: 8px; }
        ::-webkit-scrollbar-track { background: #0B1121; }
        ::-webkit-scrollbar-thumb { background: #2C3652; border-radius: 4px; }
        ::-webkit-scrollbar-thumb:hover { background: #00D2FF; }
        /* Main App & Sidebar Backgrounds */
        .stApp { background-color: #0B1121; color: #FFFFFF; }
        [data-testid="stSidebar"] { background-color: #0F1423; border-right: 1px solid #1E2638; }
        
        /* Customizing Sidebar Navigation */
        [data-testid="stSidebar"] .stRadio > div[role="radiogroup"] { gap: 8px; }
        [data-testid="stSidebar"] .stRadio > div[role="radiogroup"] > label {
            background-color: transparent;
            padding: 12px 15px;
            border-radius: 8px;
            margin-bottom: 2px;
            transition: all 0.2s ease;
            cursor: pointer;
            color: #9CA3AF;
        }
        [data-testid="stSidebar"] .stRadio > div[role="radiogroup"] > label:hover {
            background-color: #1A2235;
            color: #FFFFFF;
        }
        [data-testid="stSidebar"] .stRadio > div[role="radiogroup"] > label[data-checked="true"] {
            background: rgba(0, 210, 255, 0.1) !important;
            border-left: 4px solid #00D2FF;
            color: #00D2FF !important;
            border-radius: 4px 8px 8px 4px;
        }
        
        /* Modern Header Banners */
        .enterprise-header {
            padding: 1.8rem;
            background: linear-gradient(135deg, #151B2E 0%, #0B1121 100%);
            border-radius: 12px;
            border: 1px solid #1E2638;
            border-left: 6px solid #00D2FF;
            margin-bottom: 2rem;
            box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.5);
        }
        .enterprise-header h1 { margin: 0; padding: 0; font-size: 2.2rem; color: #FFFFFF; font-weight: 800; letter-spacing: -0.5px; }
        .enterprise-header p { margin: 8px 0 0 0; color: #00D2FF; font-size: 1rem; text-transform: uppercase; letter-spacing: 1.5px; font-weight: 700; }
        /* Metric Cards (KPIs) Transformation */
        [data-testid="stMetric"] {
            background-color: #151B2E;
            border: 1px solid #1E2638;
            padding: 20px 24px;
            border-radius: 12px;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.3);
            transition: transform 0.2s ease, box-shadow 0.2s ease, border-color 0.2s ease;
        }
        [data-testid="stMetric"]:hover {
            transform: translateY(-4px);
            box-shadow: 0 10px 20px -3px rgba(0, 210, 255, 0.1);
            border-color: #00D2FF;
        }
        [data-testid="stMetricLabel"] { color: #9CA3AF; font-weight: 600; text-transform: uppercase; font-size: 0.85rem; letter-spacing: 0.5px; }
        [data-testid="stMetricValue"] { color: #FFFFFF; font-weight: 800; font-size: 2.2rem; }
        [data-testid="stMetricDelta"] { font-weight: 700; color: #00FF9D; }
        /* Beautiful Buttons */
        .stButton>button {
            background: linear-gradient(135deg, #00D2FF 0%, #3A86FF 100%);
            color: #FFFFFF;
            border: none;
            border-radius: 8px;
            padding: 0.6rem 1.2rem;
            font-weight: 600;
            letter-spacing: 0.5px;
            transition: all 0.3s ease;
            box-shadow: 0 4px 10px rgba(0, 210, 255, 0.3);
        }
        .stButton>button:hover { 
            box-shadow: 0 6px 15px rgba(0, 210, 255, 0.5); 
            transform: translateY(-2px); 
            filter: brightness(1.1);
        }
        /* Dataframes & Tables */
        .stDataFrame { border-radius: 12px; overflow: hidden; border: 1px solid #1E2638; box-shadow: 0 4px 6px rgba(0,0,0,0.2); }
        
        /* Expander & Forms */
        .stExpander { background-color: #151B2E; border: 1px solid #1E2638 !important; border-radius: 10px !important; }
        [data-testid="stForm"] { background-color: #0F1423; border: 1px solid #1E2638; border-radius: 16px; padding: 25px; box-shadow: 0 10px 20px rgba(0,0,0,0.4); }
        
        /* Tabs Styled as Modern Pills */
        .stTabs [data-baseweb="tab-list"] { background-color: transparent; gap: 8px; margin-bottom: 15px; }
        .stTabs [data-baseweb="tab"] { 
            background-color: #151B2E; 
            border-radius: 20px !important; 
            border: 1px solid #1E2638; 
            color: #9CA3AF; 
            padding: 8px 20px !important;
            height: auto;
        }
        .stTabs [aria-selected="true"] { 
            background-color: #00D2FF !important; 
            color: #0B1121 !important; 
            border-color: #00D2FF !important; 
            font-weight: 700;
        }
        
        /* Form Inputs Customization */
        .stTextInput input, .stSelectbox div[data-baseweb="select"], .stDateInput input {
            background-color: #1A2235 !important;
            border: 1px solid #2C3652 !important;
            color: #FFFFFF !important;
            border-radius: 8px !important;
        }
        .stTextArea textarea {
            background-color: #1A2235 !important;
            border: 1px solid #2C3652 !important;
            color: #FFFFFF !important;
            border-radius: 8px !important;
        }

        /* Report-card panel styles, defined once globally so Machine Explorer,
           AI Assistant, etc. render correctly even before render_visual_report_card runs. */
        .report-box { background-color: #151B2E; border: 1px solid #1E2638; border-radius: 16px; padding: 24px; color: #FFFFFF; margin-bottom: 25px; box-shadow: 0 10px 20px rgba(0,0,0,0.3); }
        .report-header { display: flex; justify-content: space-between; align-items: center; border-bottom: 1px solid #1E2638; padding-bottom: 15px; margin-bottom: 15px; }
        .meta-bar { display: flex; justify-content: space-between; background-color: #0F1423; padding: 12px 18px; border-radius: 10px; font-size: 0.9rem; margin-bottom: 20px; border: 1px solid #1E2638; }
        .card-panel { background-color: #0F1423; border: 1px solid #1E2638; border-radius: 12px; padding: 20px; height: 100%; box-shadow: 0 4px 6px rgba(0,0,0,0.2); }
        .badge-tag { padding: 4px 10px; border-radius: 6px; font-size: 0.75rem; font-weight: bold; display: inline-block; }
        .kpi-title { font-size: 0.8rem; color: #9CA3AF; font-weight: 700; text-transform: uppercase; margin-bottom: 6px; letter-spacing: 0.5px; }
        .kpi-val { font-size: 1.6rem; font-weight: 800; color: #FFFFFF; }
        .metric-card { background-color: #151B2E; border-radius: 10px; padding: 15px; text-align: left; border: 1px solid #1E2638; margin-bottom: 10px; transition: transform 0.2s; }
        .metric-card:hover { transform: translateY(-2px); border-color: #00D2FF; }
        .ai-report-box strong { color: #00D2FF; }
    </style>
    """, unsafe_allow_html=True)

inject_enterprise_ui()

def create_header(title, subtitle):
    st.markdown(f"""
        <div class="enterprise-header">
            <h1>{title}</h1>
            <p>{subtitle}</p>
        </div>
    """, unsafe_allow_html=True)

# Plotly Global Theme Fix
plotly_layout_defaults = dict(
    plot_bgcolor='rgba(0,0,0,0)',
    paper_bgcolor='rgba(0,0,0,0)',
    font_color='#9CA3AF',
    title_font_color='#FFFFFF',
    margin=dict(l=20, r=20, t=40, b=20)
)

# Vibrant Tech Chart Colors (Cyan, Blue, Purple, Pink, Lime)
custom_colors = ['#00D2FF', '#3A86FF', '#8A2BE2', '#FF007A', '#00FF9D']

@st.cache_data
def load_data():
    return pd.read_csv(DATA_FILE)

try:
    df = load_data()
except Exception as e:
    st.error(f"Error loading 'ai4i2020.csv'. Please place the file in your working directory. Details: {e}")
    st.stop()

# ==============================================================================
# SESSION STATE INITIALIZATION
# ==============================================================================
if 'logged_in' not in st.session_state: st.session_state['logged_in'] = False
if 'current_user' not in st.session_state: st.session_state['current_user'] = None
if 'current_role' not in st.session_state: st.session_state['current_role'] = None
if 'nav_mode' not in st.session_state: st.session_state['nav_mode'] = "Dashboard"
if 'auto_run_ai' not in st.session_state: st.session_state['auto_run_ai'] = False
if 'last_redirected_id' not in st.session_state: st.session_state['last_redirected_id'] = None
if 'prefilled_wo' not in st.session_state:
    st.session_state['prefilled_wo'] = {'product_id': df['Product ID'].iloc[0], 'maint_type': 'Preventive Maintenance', 'priority': 'Medium'}
if 'pm_auto_fill' not in st.session_state:
    st.session_state['pm_auto_fill'] = {'machine_id': df['Product ID'].iloc[0], 'frequency': 'Monthly', 'checklist': "1. Visual Inspection\n2. Lubrication\n3. Sensor Calibration\n4. Tool Wear Check"}
# NEW: single canonical "currently active machine" used to keep the target
# machine ID consistent across Machine Explorer, AI Maintenance Assistant,
# the Preventive Maintenance "AI Recommendations" tab, and the "Create Schedule" tab.
if 'target_machine_id' not in st.session_state:
    st.session_state['target_machine_id'] = df['Product ID'].iloc[0]

def sync_target_machine(product_id):
    """
    Keeps the target machine ID synchronized across modules so the user only
    has to pick a machine once. Updates:
      - the canonical target_machine_id
      - selected_machine (used by AI Maintenance Assistant / Machine Explorer)
      - prefilled_wo['product_id'] (used by Work Order Creation)
      - pm_auto_fill['machine_id'] (used by PM Scheduler > Create Schedule)
    Note: because Streamlit executes all tab bodies on every rerun in source
    order, a change made in a later tab (e.g. AI Recommendations) will be
    reflected in an earlier tab (e.g. Create Schedule) starting on the NEXT
    rerun, not instantly within the same script pass.
    """
    st.session_state['target_machine_id'] = product_id
    match = df[df['Product ID'] == product_id]
    if not match.empty:
        st.session_state['selected_machine'] = match.iloc[0].to_dict()
    st.session_state['prefilled_wo']['product_id'] = product_id
    st.session_state['pm_auto_fill']['machine_id'] = product_id

def extract_work_order_defaults(m_dict, report_text):
    report_text_lower = report_text.lower() if report_text else ""
    priority = "Medium"
    if m_dict.get('Machine failure', 0) == 1 or "critical" in report_text_lower or "high" in report_text_lower:
        priority = "High"
    elif "low" in report_text_lower:
        priority = "Low"
    maint_type = "Preventive Maintenance"
    if priority == "High" or "emergency" in report_text_lower or "repair" in report_text_lower:
        maint_type = "Emergency Repair"
    elif "inspection" in report_text_lower:
        maint_type = "Inspection"
    elif "calibration" in report_text_lower:
        maint_type = "Calibration"
    return {'product_id': m_dict.get('Product ID', df['Product ID'].iloc[0]), 'maint_type': maint_type, 'priority': priority}

# ==============================================================================
# HELPER FUNCTIONS
# ==============================================================================
def clean_text_for_pdf(text):
    if not text: return ""
    text = str(text)
    text = re.sub(r'\*+', '', text)
    text = re.sub(r'#+\s*', '', text)
    text = re.sub(r'`+', '', text)
    text = re.sub(r'_{2,}', '', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    replacements = {
        '🚨': '[CRITICAL]', '✅': '[OK]', '⚠️': '[WARNING]', '🔴': '[FAIL]',
        '🟢': '[NORMAL]', '🟡': '[UPCOMING]', '⚙️': '', '🤖': '', '📋': '', '🔬': '',
        '⚡': '', '🛠️': '', '📝': '', '🚀': '', '📄': '', '🔨': '', '📅': '', '🗓️': '', '➕': '', '📜': '',
        '•': '-', '–': '-', '—': '-', '“': '"', '”': '"', '‘': "'", '’': "'",
        '\u2022': '-', '\u2013': '-', '\u2014': '-', '\u201c': '"', '\u201d': '"',
        '\u2018': "'", '\u2019': "'"
    }
    for k, v in replacements.items(): text = text.replace(k, v)
    return text.encode('latin-1', 'ignore').decode('latin-1')

def is_heading_line(line):
    """
    Detects a markdown-style bold heading occupying a whole line, e.g.
    '**Executive Summary**'. Returns the inner heading text if matched,
    otherwise None.
    """
    m = re.match(r'^\*\*(.+?)\*\*\s*$', line.strip())
    return m.group(1).strip() if m else None

def render_pdf_report_body(pdf, text):
    """
    Writes `text` into an fpdf2 PDF, rendering whole-line '**Heading**'
    markers as bold text and everything else as normal body text.
    """
    for raw_line in text.split('\n'):
        stripped = raw_line.strip()
        heading_text = is_heading_line(raw_line)
        if heading_text:
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(0, 0, 0)
            pdf.multi_cell(0, 7, clean_text_for_pdf(heading_text))
            pdf.set_font("Helvetica", "", 10)
        elif stripped == "":
            pdf.ln(3)
        else:
            pdf.multi_cell(0, 6, clean_text_for_pdf(raw_line))

def markdown_bold_to_html(escaped_text):
    """
    Converts '**text**' markers into <strong> tags. Expects text that has
    ALREADY been through html.escape() so this only needs to handle the
    literal asterisk markers, not raw HTML injection.
    """
    return re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', escaped_text)

def get_next_id(df, col_name, prefix):
    if df.empty: return f"{prefix}-0001"
    try:
        nums = df[col_name].str.extract(f"{prefix}-(\\d+)")[0].dropna().astype(int)
        if not nums.empty: return f"{prefix}-{nums.max() + 1:04d}"
    except Exception as e:
        print(f"[get_next_id] fallback used due to: {e}")
    return f"{prefix}-{len(df) + 1:04d}"

def get_fpdf_byte_stream(pdf_obj):
    try:
        pdf_out = pdf_obj.output(dest='S')
        if isinstance(pdf_out, str): return pdf_out.encode('latin-1', 'replace')
        return bytes(pdf_out)
    except TypeError:
        pdf_out = pdf_obj.output()
        if isinstance(pdf_out, (bytes, bytearray)): return bytes(pdf_out)
        return str(pdf_out).encode('latin-1', 'replace')

def extract_json_from_llm(text):
    try:
        start = text.find('{')
        end = text.rfind('}') + 1
        if start != -1 and end != 0: return json.loads(text[start:end])
        return json.loads(text)
    except Exception as e:
        print(f"[extract_json_from_llm] could not parse JSON: {e}")
        return None

def get_fallback_json(m_row, raw_text=""):
    is_failed = m_row.get('Machine failure', 0) == 1
    health = 45 if is_failed else max(60, int(100 - (m_row.get('Tool wear [min]', 0) / 240.0 * 20)))
    return {
        "machine_health": health, "overall_status": "Critical" if is_failed else "Healthy",
        "failure_probability": "85%" if is_failed else "14%",
        "remaining_useful_life": f"{max(0, 240 - int(m_row.get('Tool wear [min]', 0)))} Hours",
        "estimated_cost": "₹4,500" if is_failed else "₹1,200", "downtime_risk": "High" if is_failed else "Low",
        "priority": "High" if is_failed else "Medium",
        "root_causes": [{"title": "Analysis Result", "confidence": 100, "severity": "High" if is_failed else "Low", "description": raw_text[:200] + "..." if raw_text else "No automated root cause parsed."}],
        "recommendations": ["Perform standard visual inspection", "Check operational telemetry"],
        "maintenance_schedule": {"Daily": "Visual Inspection", "Monthly": "Tool Wear Check", "Quarterly": "Sensor Calibration", "Yearly": "Complete Overhaul"},
        "executive_summary": raw_text if raw_text and len(raw_text) < 300 else "Standard operating procedure recommended based on baseline metrics."
    }

# ==============================================================================
# DATABASE & AUTHENTICATION MANAGEMENT
# ==============================================================================
def init_db():
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute('''CREATE TABLE IF NOT EXISTS users (username TEXT PRIMARY KEY, password TEXT NOT NULL, role TEXT NOT NULL)''')
    cursor.execute("SELECT COUNT(*) FROM users")
    if cursor.fetchone()[0] == 0:
        admin_pw = hashlib.sha256(b"admin123").hexdigest()
        tech_pw = hashlib.sha256(b"tech123").hexdigest()
        cursor.executemany("INSERT INTO users VALUES (?, ?, ?)", [("admin", admin_pw, "Admin"), ("david", tech_pw, "Technician")])
    cursor.execute('''CREATE TABLE IF NOT EXISTS work_orders (
        work_order_id TEXT PRIMARY KEY, product_id TEXT NOT NULL, maintenance_type TEXT NOT NULL,
        priority TEXT NOT NULL, assigned_technician TEXT NOT NULL, date_assigned TEXT NOT NULL,
        status TEXT NOT NULL, ai_report TEXT)''')
    try: cursor.execute('ALTER TABLE work_orders ADD COLUMN ai_report TEXT')
    except sqlite3.OperationalError: pass
    cursor.execute('''CREATE TABLE IF NOT EXISTS pm_schedules (
        schedule_id TEXT PRIMARY KEY, product_id TEXT NOT NULL, frequency TEXT NOT NULL,
        next_due_date TEXT NOT NULL, technician TEXT NOT NULL, checklist TEXT)''')
    conn.commit()
    conn.close()

def verify_login(username, password):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    hashed_pw = hashlib.sha256(password.encode()).hexdigest()
    cursor.execute("SELECT role FROM users WHERE username=? AND password=?", (username, hashed_pw))
    result = cursor.fetchone()
    conn.close()
    if result: return result[0]
    return None

def add_work_order(wo_id, product_id, maint_type, priority, technician, date_assigned, status, ai_report=""):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute('''INSERT INTO work_orders (work_order_id, product_id, maintenance_type, priority, assigned_technician, date_assigned, status, ai_report)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)''', (wo_id, product_id, maint_type, priority, technician, str(date_assigned), status, ai_report))
    conn.commit()
    conn.close()

def get_all_work_orders():
    conn = sqlite3.connect(DB_FILE)
    df_wo = pd.read_sql_query("SELECT * FROM work_orders", conn)
    conn.close()
    if not df_wo.empty: df_wo.columns = ["Work Order ID", "Product ID", "Maintenance Type", "Priority", "Assigned Technician", "Date Assigned", "Status", "AI Report"]
    else: df_wo = pd.DataFrame(columns=["Work Order ID", "Product ID", "Maintenance Type", "Priority", "Assigned Technician", "Date Assigned", "Status", "AI Report"])
    return df_wo

def update_work_order_status(wo_id, new_status):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("UPDATE work_orders SET status = ? WHERE work_order_id = ?", (new_status, wo_id))
    conn.commit()
    conn.close()

def delete_work_order(wo_id):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM work_orders WHERE work_order_id = ?", (wo_id,))
    conn.commit()
    conn.close()

def add_pm_schedule(sched_id, product_id, frequency, next_due_date, technician, checklist):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute('''INSERT INTO pm_schedules (schedule_id, product_id, frequency, next_due_date, technician, checklist)
        VALUES (?, ?, ?, ?, ?, ?)''', (sched_id, product_id, frequency, str(next_due_date), technician, checklist))
    conn.commit()
    conn.close()

def get_all_pm_schedules():
    conn = sqlite3.connect(DB_FILE)
    df_pm = pd.read_sql_query("SELECT * FROM pm_schedules", conn)
    conn.close()
    if not df_pm.empty: df_pm.columns = ["Schedule ID", "Product ID", "Frequency", "Next Due Date", "Technician", "Checklist"]
    else: df_pm = pd.DataFrame(columns=["Schedule ID", "Product ID", "Frequency", "Next Due Date", "Technician", "Checklist"])
    return df_pm

def update_pm_schedule_date(sched_id, new_date):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("UPDATE pm_schedules SET next_due_date = ? WHERE schedule_id = ?", (str(new_date), sched_id))
    conn.commit()
    conn.close()

def delete_pm_schedule(sched_id):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM pm_schedules WHERE schedule_id = ?", (sched_id,))
    conn.commit()
    conn.close()

init_db()

# ==============================================================================
# SECURE LOGIN SCREEN
# ==============================================================================
if not st.session_state['logged_in']:
    st.markdown("<br><br><div style='text-align: center;'><h1 style='color: #00D2FF; font-size: 3.5rem; margin-bottom: 0;'>Agentic AI for Smart Facility Operations and Optimization</h1><p style='color: #9CA3AF; font-size: 1.2rem; letter-spacing: 2px; text-transform: uppercase;'>Enterprise Authentication</p></div><br>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 1.2, 1])
    with col2:
        with st.form("login_form"):
            st.markdown("<h3 style='color: #FFFFFF; margin-top: 0;'>Secure Login</h3>", unsafe_allow_html=True)
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            submit = st.form_submit_button("Authenticate")
            
            if submit:
                role = verify_login(username, password)
                if role:
                    st.session_state['logged_in'] = True
                    st.session_state['current_user'] = username
                    st.session_state['current_role'] = role
                    st.success(f"✅ Process Completed: Logged in successfully as **{username}** ({role}).")
                    st.rerun()
                else:
                    st.error("Invalid username or password.")
                    
        st.markdown("<br>", unsafe_allow_html=True)
        st.info("💡 **Demo Accounts:**\n\n**Admin Role:** `admin` / `admin123`\n\n**Technician Role:** `david` / `tech123`")
    
    st.stop() 

# ==============================================================================
# BULLETPROOF DOCUMENT GENERATORS
# ==============================================================================
def generate_pdf_work_order(wo_data, machine_df):
    wo_id = wo_data.get("Work Order ID", "WO-0000")
    prod_id = wo_data.get("Product ID", "N/A")
    ai_raw = str(wo_data.get("AI Report", ""))
    if not ai_raw or ai_raw.strip() in ["", "None", "nan"]:
        ai_raw = "No custom AI Maintenance Report attached. Proceed with standard baseline inspection."
        
    current_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    try:
        from weasyprint import HTML
        ai_raw_escaped = html.escape(ai_raw)
        ai_raw_escaped = markdown_bold_to_html(ai_raw_escaped)
        ai_raw_escaped = ai_raw_escaped.replace('\n', '<br>')
            
        html_content = f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="UTF-8"><style>
    @page {{ size: A4; margin: 15mm; background-color: #0B1121; }}
    body {{ font-family: 'Segoe UI', Helvetica, sans-serif; background: #0B1121; color: #FFFFFF; margin: 0; padding: 0; font-size: 10pt; }}
    .header {{ border-bottom: 2px solid #1E2638; padding-bottom: 15px; margin-bottom: 20px; }}
    .title {{ font-size: 18pt; font-weight: 800; color: #00D2FF; margin-bottom: 5px; }}
    .subtitle {{ font-size: 10pt; color: #9CA3AF; margin-bottom: 3px; }}
    .h1 {{ font-size: 14pt; color: #00D2FF; font-weight: bold; margin-top: 20px; margin-bottom: 10px; border-bottom: 1px solid #1E2638; padding-bottom: 5px; }}
    .summary {{ background: #151B2E; padding: 15px; border-radius: 8px; font-size: 10pt; color: #FFFFFF; border: 1px solid #1E2638; line-height: 1.5; white-space: pre-wrap; }}
    .summary strong {{ font-weight: 800; color: #00D2FF; display: inline-block; margin-top: 8px; }}
</style></head><body>
    <div class="header">
        <div class="title">FacilityOps AI: Predictive Maintenance Report</div>
        <div class="subtitle">Work Order: {html.escape(str(wo_id))} | Machine ID: {html.escape(str(prod_id))}</div>
        <div class="subtitle">Generated: {current_time}</div>
    </div>
    <div class="h1">AI Diagnostic Analysis & Action Plan</div>
    <div class="summary">{ai_raw_escaped}</div>
</body></html>"""
        pdf_out = HTML(string=html_content).write_pdf()
        if not pdf_out: raise ValueError("WeasyPrint returned empty bytes")
        return pdf_out
    except Exception as e:
        from fpdf import FPDF
        pdf = FPDF(unit='mm', format='A4')
        pdf.set_margins(15, 15, 15)
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        
        safe_wo_id = clean_text_for_pdf(wo_id)
        safe_prod_id = clean_text_for_pdf(prod_id)
        
        # Exact RGB mappings for the new deep space schema
        pdf.set_fill_color(11, 17, 33) # #0B1121
        pdf.rect(0, 0, 210, 40, 'F')
        pdf.set_y(12)
        pdf.set_font("Helvetica", "B", 16)
        pdf.set_text_color(0, 210, 255) # #00D2FF
        pdf.cell(0, 8, "FacilityOps AI: Predictive Maintenance Report", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(156, 163, 175) # #9CA3AF
        pdf.cell(0, 6, f"Work Order: {safe_wo_id} | Machine ID: {safe_prod_id}", ln=True)
        pdf.cell(0, 6, f"Generated: {current_time}", ln=True)
        
        pdf.set_y(45)
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(0, 210, 255) # #00D2FF
        pdf.cell(0, 10, "AI Diagnostic Analysis & Action Plan", ln=True)
        
        pdf.set_text_color(0, 0, 0) # Base text color
        pdf.set_font("Helvetica", "", 10)
        render_pdf_report_body(pdf, ai_raw)
        
        return get_fpdf_byte_stream(pdf)

def generate_docx_work_order(wo_data, machine_df):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        ai_raw = str(wo_data.get("AI Report", ""))
        if not ai_raw or ai_raw.strip() in ["", "None", "nan"]:
            ai_raw = "No custom AI Maintenance Report attached."
        head = doc.add_heading("FacilityOps AI: Predictive Maintenance Report", 0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Work Order: {wo_data.get('Work Order ID', 'N/A')} | Machine ID: {wo_data.get('Product ID', 'N/A')}")
        doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_heading("AI Diagnostic Analysis & Action Plan", level=1)
        for line in ai_raw.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue
            heading_text = is_heading_line(line)
            if heading_text:
                doc.add_heading(heading_text, level=2)
            else:
                doc.add_paragraph(stripped)
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io.getvalue()
    except Exception as e:
        output = f"FACILITYOPS WORK ORDER\n{'='*30}\n"
        for col, val in wo_data.items(): output += f"{col}: {val}\n"
        return output.encode('utf-8')

# ==============================================================================
# VISUAL DASHBOARD REPORT CARD RENDERER
# ==============================================================================
def render_visual_report_card(wo_data, machine_df):
    prod_id = html.escape(str(wo_data.get("Product ID", "M14860")))
    wo_id = html.escape(str(wo_data.get("Work Order ID", "WO-0001")))
    tech_name = html.escape(str(wo_data.get("Assigned Technician", "David")))
    
    raw_prod_id = str(wo_data.get("Product ID", "M14860"))
    m_match = machine_df[machine_df['Product ID'] == raw_prod_id] if not machine_df.empty else pd.DataFrame()
    
    if not m_match.empty:
        m_row = m_match.iloc[0]
        air_temp, proc_temp, rpm, torque, tool_wear, is_failed, m_type = (
            m_row.get('Air temperature [K]', 298.1), m_row.get('Process temperature [K]', 308.6),
            m_row.get('Rotational speed [rpm]', 1551), m_row.get('Torque [Nm]', 42.8),
            m_row.get('Tool wear [min]', 0), m_row.get('Machine failure', 0) == 1, m_row.get('Type', 'M')
        )
    else:
        m_row = {}
        air_temp, proc_temp, rpm, torque, tool_wear, is_failed, m_type = 298.1, 308.6, 1551, 42.8, 0, False, 'M'
    m_type = html.escape(str(m_type))
    ai_raw = str(wo_data.get("AI Report", ""))
    
    ai_data = extract_json_from_llm(ai_raw)
    if not ai_data: ai_data = get_fallback_json(m_row, ai_raw)
    health_score = ai_data.get("machine_health", 100)
    status_text = str(ai_data.get("overall_status", "Healthy")).upper()
    badge_bg = "#00FF9D" if status_text == "HEALTHY" else "#FF007A"
    badge_fg = "#0B1121"

    gen_time = datetime.datetime.now().strftime('%d %b %Y %I:%M %p')
    st.markdown(f"""
    <div class="report-box">
        <div class="report-header">
            <div>
                <h2 style="margin:0; color:#00D2FF; font-weight:800;">FacilityOps AI</h2>
                <span style="font-size:0.85rem; color:#9CA3AF; text-transform:uppercase; letter-spacing:1px; font-weight:600;">Predictive Maintenance Dashboard</span>
            </div>
            <div style="text-align:right;">
                <h3 style="margin:0; color:#FFFFFF;">WORK ORDER SUMMARY</h3>
                <span class="badge-tag" style="background-color:{badge_bg}; color:{badge_fg}; margin-top:4px;">{html.escape(status_text)}</span>
            </div>
        </div>
        <div class="meta-bar">
            <div><span style="color:#9CA3AF;">Work Order ID:</span> <b style="color:#FFFFFF;">{wo_id}</b></div>
            <div><span style="color:#9CA3AF;">Machine ID:</span> <b style="color:#00D2FF;">{prod_id}</b></div>
            <div><span style="color:#9CA3AF;">Variant:</span> <b style="color:#FFFFFF;">Type {m_type}</b></div>
            <div><span style="color:#9CA3AF;">Technician:</span> <b style="color:#FFFFFF;">{tech_name}</b></div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    r1_col1, r1_col2 = st.columns([1, 2.5])
    with r1_col1:
        st.markdown('<div class="card-panel">', unsafe_allow_html=True)
        st.markdown('<div class="kpi-title">MACHINE HEALTH SCORE</div>', unsafe_allow_html=True)
        fig_g = go.Figure(go.Indicator(
            mode = "gauge+number", value = health_score,
            number = {'suffix': "%", 'font': {'color': '#FFFFFF', 'size': 36}},
            gauge = {
                'axis': {'range': [0, 100], 'tickcolor': "#1E2638"},
                'bar': {'color': badge_bg}, 'bgcolor': "#151B2E", 'bordercolor': "#1E2638",
                'steps': [{'range': [0, 50], 'color': 'rgba(255, 0, 122, 0.2)'}, {'range': [50, 80], 'color': 'rgba(58, 134, 255, 0.2)'}, {'range': [80, 100], 'color': 'rgba(0, 255, 157, 0.2)'}]
            }
        ))
        fig_g.update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', height=190, margin=dict(l=10, r=10, t=10, b=10))
        st.plotly_chart(fig_g, use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)
    with r1_col2:
        st.markdown('<div class="card-panel"><div class="kpi-title" style="margin-bottom:12px;">LIVE TELEMETRY SUMMARY</div>', unsafe_allow_html=True)
        s1, s2, s3, s4 = st.columns(4)
        s1.markdown(f'<div class="metric-card"><div class="kpi-title">Air Temp</div><div class="kpi-val">{air_temp} <span style="font-size:0.8rem;color:#9CA3AF">K</span></div></div>', unsafe_allow_html=True)
        s2.markdown(f'<div class="metric-card"><div class="kpi-title">Process Temp</div><div class="kpi-val">{proc_temp} <span style="font-size:0.8rem;color:#9CA3AF">K</span></div></div>', unsafe_allow_html=True)
        s3.markdown(f'<div class="metric-card"><div class="kpi-title">RPM</div><div class="kpi-val">{rpm}</div></div>', unsafe_allow_html=True)
        s4.markdown(f'<div class="metric-card"><div class="kpi-title">Torque</div><div class="kpi-val">{torque} <span style="font-size:0.8rem;color:#9CA3AF">Nm</span></div></div>', unsafe_allow_html=True)
        st.markdown(f"""<div style="display:flex; justify-content:space-between; align-items:center; background:#151B2E; padding:12px 18px; border-radius:10px; border:1px solid #1E2638;">
            <div><span class="kpi-title">Tool Wear Status:</span> <b style="color:#00D2FF; font-size:1.2rem; margin-left:10px;">{tool_wear} min</b> / 240 min</div>
            <div><progress value="{tool_wear}" max="240" style="width: 250px; height: 14px; accent-color: #3A86FF; border-radius: 6px;"></progress></div>
        </div></div>""", unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div class="card-panel">', unsafe_allow_html=True)
    st.markdown('<div class="kpi-title" style="margin-bottom:10px; color:#00D2FF;">AI DIAGNOSTIC REPORT</div>', unsafe_allow_html=True)
    
    if not ai_raw or ai_raw.strip() in ["", "None", "nan"]:
        ai_raw = f"Machine {raw_prod_id} is operating in a manageable condition. Preventive inspection of cooling channels is recommended."

    ai_raw_safe = html.escape(ai_raw)
    ai_raw_safe = markdown_bold_to_html(ai_raw_safe)
    ai_raw_safe = ai_raw_safe.replace('\n', '<br>')
    st.markdown(f'<div class="ai-report-box" style="font-size:1rem; color:#FFFFFF; line-height:1.6; white-space:pre-wrap; background:#151B2E; padding: 20px; border-radius: 10px; border: 1px solid #1E2638;">{ai_raw_safe}</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

# ==============================================================================
# SIDEBAR NAVIGATION & RBAC
# ==============================================================================
st.sidebar.title("🛠️ Operations Hub")

# Role-Based Routing Array
all_modules = [
    "Dashboard", 
    "Data Analysis (EDA)", 
    "Machine Explorer", 
    "AI Maintenance Assistant",
    "Work Order Creation",
    "Work Order Management",
    "Preventive Maintenance Scheduler"
]
tech_modules = [
    "Dashboard", 
    "Machine Explorer", 
    "Work Order Management"
]
available_modules = all_modules if st.session_state['current_role'] == 'Admin' else tech_modules
app_mode = st.sidebar.radio("Navigation:", available_modules, key='nav_mode')

# Logout System
st.sidebar.markdown("---")
st.sidebar.markdown(f"<span style='color:#9CA3AF;'>👤 Logged in as:</span> **<span style='color:#00D2FF;'>{html.escape(str(st.session_state['current_user']))}</span>** ({html.escape(str(st.session_state['current_role']))})", unsafe_allow_html=True)
if st.sidebar.button("Logout", key="logout_btn"):
    st.session_state['logged_in'] = False
    st.session_state['current_user'] = None
    st.session_state['current_role'] = None
    st.rerun()

if app_mode == "Dashboard":
    st.sidebar.markdown("---")
    st.sidebar.subheader("🔍 Global Filters")
    selected_status = st.sidebar.selectbox("Machine Status", ["All", "Healthy", "Failed"])
    selected_type = st.sidebar.selectbox("Machine Type", ["All"] + list(df['Type'].unique()))
    
    filtered_df = df.copy()
    if selected_status == "Healthy": filtered_df = filtered_df[filtered_df['Machine failure'] == 0]
    elif selected_status == "Failed": filtered_df = filtered_df[filtered_df['Machine failure'] == 1]
    if selected_type != "All": filtered_df = filtered_df[filtered_df['Type'] == selected_type]

# ==============================================================================
# DASHBOARD
# ==============================================================================
if app_mode == "Dashboard":
    create_header("Agentic AI for Smart Facility Operations and Optimization", "Real-time Predictive Maintenance Telemetry")
    
    k1, k2, k3, k4 = st.columns(4)
    tot = len(filtered_df)
    failed_cnt = int(filtered_df['Machine failure'].sum())
    healthy_cnt = tot - failed_cnt
    avg_rpm = filtered_df['Rotational speed [rpm]'].mean() if tot > 0 else 0
    k1.metric("Total Machines", f"{tot:,}")
    k2.metric("Healthy Units", f"{healthy_cnt:,}")
    k3.metric("Failed Units", f"{failed_cnt:,}", delta="-Action Required" if failed_cnt > 0 else None, delta_color="inverse")
    k4.metric("Avg Rotational Speed", f"{avg_rpm:.1f} RPM")
    
    k5, k6, k7, k8 = st.columns(4)
    avg_torq = filtered_df['Torque [Nm]'].mean() if tot > 0 else 0
    avg_wear = filtered_df['Tool wear [min]'].mean() if tot > 0 else 0
    avg_temp = filtered_df['Air temperature [K]'].mean() if tot > 0 else 0
    k5.metric("Avg Torque", f"{avg_torq:.2f} Nm")
    k6.metric("Avg Tool Wear", f"{avg_wear:.1f} min")
    k7.metric("Avg Air Temp", f"{avg_temp:.1f} K")
    k8.metric("System Health State", "Active Monitoring")
    
    st.markdown("---")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("<h4 style='color:#FFFFFF;'>Machine Failure Mode Distribution</h4>", unsafe_allow_html=True)
        fail_cols = ['TWF', 'HDF', 'PWF', 'OSF', 'RNF']
        avail_fails = [c for c in fail_cols if c in filtered_df.columns]
        if avail_fails and filtered_df['Machine failure'].sum() > 0:
            f_sums = filtered_df[avail_fails].sum()
            fig_fail = px.bar(x=f_sums.index, y=f_sums.values, labels={'x':'Failure Type', 'y':'Count'}, color=f_sums.index, color_discrete_sequence=custom_colors)
            fig_fail.update_layout(**plotly_layout_defaults)
            st.plotly_chart(fig_fail, use_container_width=True)
        else:
            bin_fail = filtered_df['Machine failure'].value_counts().reset_index()
            bin_fail.columns = ['Status', 'Count']
            bin_fail['Status'] = bin_fail['Status'].map({0: 'Healthy', 1: 'Failed'})
            fig_fail = px.bar(bin_fail, x='Status', y='Count', color='Status', color_discrete_map={'Healthy':'#00FF9D', 'Failed':'#FF007A'})
            fig_fail.update_layout(**plotly_layout_defaults)
            st.plotly_chart(fig_fail, use_container_width=True)
    with c2:
        st.markdown("<h4 style='color:#FFFFFF;'>Machine Variant Type Breakdown</h4>", unsafe_allow_html=True)
        t_counts = filtered_df['Type'].value_counts().reset_index()
        t_counts.columns = ['Type', 'Count']
        fig_pie = px.pie(t_counts, values='Count', names='Type', hole=0.4, color_discrete_sequence=custom_colors)
        fig_pie.update_layout(**plotly_layout_defaults)
        st.plotly_chart(fig_pie, use_container_width=True)
        
    st.markdown("---")
    c3, c4 = st.columns(2)
    with c3:
        st.markdown("<h4 style='color:#FFFFFF;'>Tool Wear Time Analysis</h4>", unsafe_allow_html=True)
        fig_hist = px.histogram(filtered_df, x="Tool wear [min]", color="Machine failure", nbins=30, color_discrete_sequence=['#3A86FF', '#FF007A'])
        fig_hist.update_layout(**plotly_layout_defaults)
        st.plotly_chart(fig_hist, use_container_width=True)
    with c4:
        st.markdown("<h4 style='color:#FFFFFF;'>Operating RPM vs Torque Correlation</h4>", unsafe_allow_html=True)
        fig_scat = px.scatter(filtered_df, x="Rotational speed [rpm]", y="Torque [Nm]", color="Machine failure", opacity=0.7, color_discrete_sequence=['#00FF9D', '#FF007A'])
        fig_scat.update_layout(**plotly_layout_defaults)
        st.plotly_chart(fig_scat, use_container_width=True)

# ==============================================================================
# EXPLORATORY DATA ANALYSIS (EDA)
# ==============================================================================
elif app_mode == "Data Analysis (EDA)":
    create_header("Exploratory Data Analysis", "Raw Dataset Telemetry & Statistical Modeling")
    
    st.markdown("<h4 style='color:#FFFFFF;'>📋 Dataset Preview</h4>", unsafe_allow_html=True)
    st.dataframe(df.head(10), use_container_width=True)
    
    col1, col2 = st.columns(2)
    with col1:
        st.info(f"**Total Records:** `{df.shape[0]}` rows | **Features:** `{df.shape[1]}` columns\n\n**Duplicated Entries:** `{df.duplicated().sum()}`")
    with col2:
        st.info(f"**Missing Values Count:** `{df.isnull().sum().sum()}`")
        
    st.markdown("---")
    st.markdown("<h4 style='color:#FFFFFF;'>📈 Statistical Summary</h4>", unsafe_allow_html=True)
    st.dataframe(df.describe(), use_container_width=True)
    
    st.markdown("---")
    st.markdown("<h4 style='color:#FFFFFF;'>🌡️ Feature Correlation Matrix</h4>", unsafe_allow_html=True)
    numeric_df = df.select_dtypes(include=[np.number])
    corr = numeric_df.corr()
    
    plt.style.use('dark_background')
    fig_heat, ax_heat = plt.subplots(figsize=(10, 5))
    fig_heat.patch.set_alpha(0.0)
    ax_heat.patch.set_alpha(0.0)
    sns.heatmap(corr, annot=True, cmap="cool", fmt=".2f", ax=ax_heat, cbar_kws={'label': 'Correlation Coefficient'})
    st.pyplot(fig_heat)

# ==============================================================================
# MACHINE EXPLORER
# ==============================================================================
elif app_mode == "Machine Explorer":
    create_header("Machine Explorer", "Individual Asset Inspection & Real-time Telemetry")
    
    product_list = ["-- Select Product ID --"] + list(df['Product ID'].dropna().unique())
    col_select, col_search = st.columns(2)
    
    with col_select:
        selected_dropdown_id = st.selectbox("Select Machine from Dropdown:", options=product_list)
        
    with col_search:
        searched_text_id = st.text_input("Or Search manually by Product ID:", "").strip()
        
    selected_id = searched_text_id if searched_text_id else (selected_dropdown_id if selected_dropdown_id != "-- Select Product ID --" else "")
    
    if selected_id:
        machine_row = df[df['Product ID'] == selected_id]
        
        if not machine_row.empty:
            m_data = machine_row.iloc[0]
            is_failed = m_data['Machine failure'] == 1
            
            if is_failed and st.session_state.get('last_redirected_id') != selected_id and st.session_state['current_role'] == 'Admin':
                st.session_state['last_redirected_id'] = selected_id
                st.session_state['auto_run_ai'] = True
                # NOTE: sync_target_machine() keeps this machine ID consistent
                # across AI Maintenance Assistant, PM AI Recommendations, and
                # PM Create Schedule.
                sync_target_machine(m_data['Product ID'])
                st.session_state['nav_mode'] = "AI Maintenance Assistant"
                st.toast(f"🚨 Faulty Machine ({selected_id}) detected! Redirecting...", icon="⚡")
                st.rerun()
                
            col_info, col_sensors = st.columns(2)
            
            with col_info:
                st.markdown("<div class='card-panel'>", unsafe_allow_html=True)
                st.markdown("<div class='kpi-title'>Machine Information</div>", unsafe_allow_html=True)
                st.write(f"**UDI Identifier:** `{m_data['UDI']}`")
                st.write(f"**Product ID:** `{m_data['Product ID']}`")
                st.write(f"**Machine Variant Type:** `{m_data['Type']}`")
                
                st.markdown("### Health Status")
                if is_failed:
                    st.error("🚨 **CRITICAL: MACHINE FAILURE DETECTED**")
                else:
                    st.success("✅ **NOMINAL: MACHINE OPERATIONAL / HEALTHY**")
                st.markdown("</div>", unsafe_allow_html=True)
                
            with col_sensors:
                st.markdown("<div class='card-panel'>", unsafe_allow_html=True)
                st.markdown("<div class='kpi-title'>Live Sensor Values</div>", unsafe_allow_html=True)
                m1, m2 = st.columns(2)
                m1.metric("Air Temperature", f"{m_data['Air temperature [K]']} K")
                m2.metric("Process Temperature", f"{m_data['Process temperature [K]']} K")
                
                m3, m4 = st.columns(2)
                m3.metric("Rotational Speed", f"{m_data['Rotational speed [rpm]']} RPM")
                m4.metric("Torque", f"{m_data['Torque [Nm]']} Nm")
                
                st.metric("Cumulative Tool Wear Time", f"{m_data['Tool wear [min]']} min")
                st.markdown("</div>", unsafe_allow_html=True)
                
            st.markdown("---")
            st.markdown("<h4 style='color:#FFFFFF;'>⚠️ Failure Analysis Diagnostics</h4>", unsafe_allow_html=True)
            fail_modes = {
                'Tool Wear Failure (TWF)': m_data.get('TWF', 0),
                'Heat Dissipation Failure (HDF)': m_data.get('HDF', 0),
                'Power Failure (PWF)': m_data.get('PWF', 0),
                'Overstrain Failure (OSF)': m_data.get('OSF', 0),
                'Random Failure (RNF)': m_data.get('RNF', 0)
            }
            
            f_cols = st.columns(5)
            for idx, (mode, val) in enumerate(fail_modes.items()):
                with f_cols[idx]:
                    if val == 1:
                        st.markdown(f"🔴 **{mode}**\n*(Triggered)*")
                    else:
                        st.markdown(f"🟢 {mode}\n*(Normal)*")
                        
            # Keep the canonical target machine in sync whenever a machine is
            # actively being viewed here, so AI Maintenance Assistant / PM
            # Scheduler default to the same machine.
            sync_target_machine(m_data['Product ID'])
        else:
            st.error("❌ Product ID not found in database. Please check the ID and try again.")
    else:
        st.info("💡 Select a Product ID from the dropdown or type one in the search box above to inspect telemetry.")

# ==============================================================================
# AI MAINTENANCE ASSISTANT
# ==============================================================================
elif app_mode == "AI Maintenance Assistant":
    create_header("AI Maintenance Assistant", "Automated Root Cause Diagnostics & Recommendation Engine")
    
    if 'selected_machine' not in st.session_state:
        st.warning("⚠️ No machine selected! Please go to **Machine Explorer** first and select a valid Product ID.")
    else:
        m_dict = st.session_state['selected_machine']
        # Keep the canonical target machine ID aligned with whatever machine
        # this report is being generated for.
        sync_target_machine(m_dict['Product ID'])
        
        c_config, c_telemetry = st.columns([1, 1])
        with c_telemetry:
            with st.expander("🔬 Active Telemetry Payload", expanded=True):
                st.json(m_dict)
        with c_config:
            with st.container():
                st.markdown("<div class='kpi-title'>⚙️ Local Model Configuration</div>", unsafe_allow_html=True)
                model_name = st.text_input("", value="llama3.2:1b", help="Change if using a different Ollama model.")
        
        st.markdown("---")
        st.markdown("<h4 style='color:#FFFFFF;'>📝 Prompt Engineering System Blueprint</h4>", unsafe_allow_html=True)
        
        # PROTECTED PROMPT — kept exactly as specified, do not modify.
        default_prompt = f"""You are an expert industrial maintenance engineer AI. Analyze this machine sensor log:
- Product ID: {m_dict['Product ID']} (Type: {m_dict['Type']})
- Air Temp: {m_dict['Air temperature [K]']} K | Process Temp: {m_dict['Process temperature [K]']} K
- RPM: {m_dict['Rotational speed [rpm]']} | Torque: {m_dict['Torque [Nm]']} Nm
- Tool Wear: {m_dict['Tool wear [min]']} minutes
- Failure Flag: {'YES (FAILED)' if m_dict['Machine failure'] == 1 else 'NO (HEALTHY)'}
Please provide a structured engineering report in plain text format. Use clear headings and bullet points. Do not return JSON.
Structure your report exactly as follows:
**Executive Summary**
[Provide a brief 1-2 sentence summary of the machine's condition and recommended next steps]
**Key Performance Indicators**
* Machine Health: [Score out of 100]%
* Overall Status: [Healthy / Warning / Critical]
* Failure Probability: [%]
* Remaining Useful Life: [Estimated Hours]
* Estimated Cost: [₹ amount]
* Downtime Risk: [Low / Medium / High]
* Priority Level: [Low / Medium / High]
**Root Cause Analysis**
* Issue: [Title of the primary issue or anomaly]
* Severity: [Low / Medium / High]
* Confidence: [%]
* Description: [Brief explanation of the metric deviations]
**Recommended Actions**
* [Actionable step 1]
* [Actionable step 2]
**Maintenance Schedule**
* Daily: [Task]
* Monthly: [Task]
* Quarterly: [Task]
* Yearly: [Task]"""
        user_prompt = st.text_area("", value=default_prompt, height=450)
        
        btn_click = st.button("🚀 Generate AI Maintenance Report (Offline Ollama)")
        auto_run = st.session_state.get('auto_run_ai', False)
        
        if btn_click or auto_run:
            st.session_state['auto_run_ai'] = False
            spinner_msg = f"🚨 Faulty Asset Detected! Auto-generating report via `{model_name}`..." if auto_run else f"Querying local Ollama core (`{model_name}`)..."
            with st.spinner(spinner_msg):
                try:
                    res = requests.post('http://localhost:11434/api/generate', json={'model': model_name, 'prompt': user_prompt, 'stream': False}, timeout=60)
                    if res.status_code == 200:
                        report_text = res.json()['response']
                        st.session_state['last_ai_report'] = report_text
                        st.session_state['prefilled_wo'] = extract_work_order_defaults(m_dict, report_text)
                        st.success(f"✅ Process Completed: AI Maintenance Report generated successfully for **{m_dict['Product ID']}** via `{model_name}`.")
                    else: st.error(f"Ollama returned status code {res.status_code}.")
                except requests.exceptions.ConnectionError:
                    st.error("⚠️ Could not connect to Ollama at http://localhost:11434. Make sure Ollama is installed and running (`ollama serve`), and that the model is pulled (`ollama pull llama3.2:1b`).")
                except requests.exceptions.Timeout:
                    st.error("⏱️ Request to Ollama timed out after 60s. The model may still be loading — try again.")
                except Exception as e:
                    st.error(f"Unexpected error while contacting Ollama: {e}")
                    
        if 'last_ai_report' in st.session_state:
            st.markdown("<h4 style='color:#00D2FF; margin-top:20px;'>📄 AI-Generated Maintenance Report</h4>", unsafe_allow_html=True)
            safe_report = html.escape(st.session_state['last_ai_report'])
            safe_report = markdown_bold_to_html(safe_report)
            safe_report = safe_report.replace('\n', '<br>')
            st.markdown(f"<div class='ai-report-box' style='background:#0F1423; padding:20px; border-radius:12px; border:1px solid #1E2638;'>{safe_report}</div>", unsafe_allow_html=True)
            st.markdown("---")
            if st.button("📋 Auto-Fill Work Order Form"):
                st.success(f"✅ Process Completed: Work Order form pre-filled with Priority **{st.session_state['prefilled_wo']['priority']}** and Maintenance Type **{st.session_state['prefilled_wo']['maint_type']}**. Switch to **Work Order Creation** in the sidebar.")

# ==============================================================================
# WORK ORDER CREATION
# ==============================================================================
elif app_mode == "Work Order Creation":
    create_header("Work Order Creation", "Generate & Store Official Maintenance Tasks")
    
    pf = st.session_state['prefilled_wo']
    all_products = list(df['Product ID'].unique())
    
    p_index = all_products.index(pf['product_id']) if pf['product_id'] in all_products else 0
    maint_options = ["Emergency Repair", "Preventive Maintenance", "Inspection", "Calibration"]
    m_index = maint_options.index(pf['maint_type']) if pf['maint_type'] in maint_options else 0
    priority_options = ["High", "Medium", "Low"]
    pr_index = priority_options.index(pf['priority']) if pf['priority'] in priority_options else 1
    current_ai_rep = st.session_state.get('last_ai_report', '')
    
    with st.form("create_work_order_db_form"):
        col_f1, col_f2, col_f3 = st.columns(3)
        with col_f1:
            target_machine = st.selectbox("Target Machine ID (Product ID):", options=all_products, index=p_index)
            maint_type = st.selectbox("Maintenance Type:", options=maint_options, index=m_index)
        with col_f2:
            priority = st.selectbox("Priority Level:", options=priority_options, index=pr_index)
            tech_options = ["David", "Alex", "Swaraj", "Aditya", "John", "Sarah", st.session_state['current_user']]
            tech_options = list(dict.fromkeys(tech_options)) 
            t_index = tech_options.index(st.session_state['current_user']) if st.session_state['current_role'] == 'Technician' else 0
            technician = st.selectbox("Assign Technician:", tech_options, index=t_index)
        with col_f3:
            date_assigned = st.date_input("Date Assigned:", value=datetime.date.today())
            order_status = st.selectbox("Initial Status:", ["Open", "In Progress", "Completed"])
            
        st.markdown("<h4 style='color:#FFFFFF; margin-top:15px;'>📄 Attached AI Diagnosis & Action Plan</h4>", unsafe_allow_html=True)
        ai_report_input = st.text_area("", value=current_ai_rep, height=250, help="This text will be embedded directly into the final PDF/DOCX.")
        
        submit_wo = st.form_submit_button("🔨 Create & Store Work Order")
        if submit_wo:
            existing_wo = get_all_work_orders()
            wo_id = get_next_id(existing_wo, 'Work Order ID', 'WO')
            add_work_order(wo_id, target_machine, maint_type, priority, technician, str(date_assigned), order_status, ai_report_input)
            st.success(f"✅ Process Completed: Work Order `{wo_id}` has been successfully created and stored for Technician **{technician}**!")

# ==============================================================================
# WORK ORDER MANAGEMENT
# ==============================================================================
elif app_mode == "Work Order Management":
    create_header("Work Order Management", "Database Overview, PDF Exports & Executive Views")
    
    wo_df = get_all_work_orders()
    
    # RBAC Filtering
    if st.session_state['current_role'] == 'Technician':
        wo_df = wo_df[wo_df['Assigned Technician'].str.lower() == str(st.session_state['current_user']).lower()]
    
    w1, w2, w3, w4 = st.columns(4)
    w1.metric("Total Work Orders", len(wo_df))
    w2.metric("Open Orders", len(wo_df[wo_df['Status'] == 'Open']) if not wo_df.empty else 0)
    w3.metric("In Progress", len(wo_df[wo_df['Status'] == 'In Progress']) if not wo_df.empty else 0)
    w4.metric("Completed", len(wo_df[wo_df['Status'] == 'Completed']) if not wo_df.empty else 0)
    
    st.markdown("---")
    st.markdown("<h4 style='color:#FFFFFF;'>🔍 Search & Filter Database</h4>", unsafe_allow_html=True)
    s_col1, s_col2, s_col3 = st.columns(3)
    with s_col1: search_query = st.text_input("Search by Work Order ID or Product ID:", "").strip()
    with s_col2: filter_status = st.selectbox("Filter by Status:", ["All", "Open", "In Progress", "Completed"])
    with s_col3: filter_priority = st.selectbox("Filter by Priority:", ["All", "High", "Medium", "Low"])
        
    filtered_wo = wo_df.copy()
    if not filtered_wo.empty:
        if search_query:
            filtered_wo = filtered_wo[
                filtered_wo['Work Order ID'].str.contains(search_query, case=False, regex=False) |
                filtered_wo['Product ID'].str.contains(search_query, case=False, regex=False)
            ]
        if filter_status != "All": filtered_wo = filtered_wo[filtered_wo['Status'] == filter_status]
        if filter_priority != "All": filtered_wo = filtered_wo[filtered_wo['Priority'] == filter_priority]
            
    st.markdown("---")
    if not filtered_wo.empty:
        display_df = filtered_wo.drop(columns=['AI Report'], errors='ignore')
        st.dataframe(display_df, use_container_width=True)
        st.markdown("---")
        
        d_col1, act_col1, act_col2 = st.columns(3)
        with d_col1:
            st.markdown("<div class='kpi-title'>📥 Export Document</div>", unsafe_allow_html=True)
            selected_dl_wo_id = st.selectbox("Select Order:", filtered_wo['Work Order ID'].tolist(), key="dl_wo_select")
            selected_row_data = filtered_wo[filtered_wo['Work Order ID'] == selected_dl_wo_id].iloc[0].to_dict()
            
            st.download_button(label="📄 Download PDF (.pdf)", data=generate_pdf_work_order(selected_row_data, df), file_name=f"{selected_dl_wo_id}_Work_Order.pdf", mime="application/pdf", use_container_width=True)
            st.download_button(label="📝 Download Word Doc (.docx)", data=generate_docx_work_order(selected_row_data, df), file_name=f"{selected_dl_wo_id}_Work_Order.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            
        with act_col1:
            st.markdown("<div class='kpi-title'>🔄 Update Status</div>", unsafe_allow_html=True)
            with st.form("update_status_form"):
                target_wo_update = st.selectbox("Work Order:", filtered_wo['Work Order ID'].tolist())
                updated_status_val = st.selectbox("New Status:", ["Open", "In Progress", "Completed"])
                if st.form_submit_button("Update Status"):
                    update_work_order_status(target_wo_update, updated_status_val)
                    st.success(f"✅ Process Completed: Work Order `{target_wo_update}` status updated to **{updated_status_val}**.")
                    st.rerun()
                    
        with act_col2:
            if st.session_state['current_role'] == 'Admin':
                st.markdown("<div class='kpi-title'>🗑️ Delete Order</div>", unsafe_allow_html=True)
                with st.form("delete_order_form"):
                    target_wo_delete = st.selectbox("Work Order:", filtered_wo['Work Order ID'].tolist())
                    confirm_delete = st.checkbox("I understand this permanently deletes this work order.")
                    if st.form_submit_button("Delete Record"):
                        if confirm_delete:
                            delete_work_order(target_wo_delete)
                            st.warning(f"🗑️ Process Completed: Work Order `{target_wo_delete}` has been permanently deleted.")
                            st.rerun()
                        else:
                            st.error("Please check the confirmation box before deleting.")
            else:
                st.markdown("<div class='kpi-title'>🗑️ Delete Order</div>", unsafe_allow_html=True)
                st.info("Admins only.")
                    
        st.markdown("---")
        render_visual_report_card(selected_row_data, df)
    else:
        st.info("💡 No matching work orders found in the database.")

# ==============================================================================
# PREVENTIVE MAINTENANCE SCHEDULER
# ==============================================================================
elif app_mode == "Preventive Maintenance Scheduler":
    create_header("Preventive Maintenance Scheduler", "Automated Task Tracking & AI Calendars")
    
    pm_df = get_all_pm_schedules()
    wo_df = get_all_work_orders()
    today = datetime.date.today()
    total_schedules = len(pm_df)
    
    overdue_count, upcoming_count = 0, 0
    if not pm_df.empty:
        pm_df['next_due_date_dt'] = pd.to_datetime(pm_df['Next Due Date']).dt.date
        overdue_count = len(pm_df[pm_df['next_due_date_dt'] < today])
        upcoming_count = len(pm_df[(pm_df['next_due_date_dt'] >= today) & (pm_df['next_due_date_dt'] <= today + datetime.timedelta(days=7))])
        
    pm_wos = len(wo_df[wo_df['Maintenance Type'] == 'Preventive Maintenance']) if not wo_df.empty else 0
    
    k1, k2, k3, k4 = st.columns(4)
    k1.metric("Active PM Schedules", total_schedules)
    k2.metric("Overdue Maintenance", overdue_count, delta="-Action Required" if overdue_count > 0 else "0", delta_color="inverse")
    k3.metric("Upcoming (7 Days)", upcoming_count)
    k4.metric("PM Work Orders Generated", pm_wos)
    st.markdown("---")
    
    tab1, tab2, tab3 = st.tabs(["📅 Calendar & Tracking", "➕ Create Schedule", "🤖 AI Recommendations"])
    
    with tab1:
        st.markdown("<h4 style='color:#FFFFFF;'>🗓️ Maintenance Calendar & Status</h4>", unsafe_allow_html=True)
        if not pm_df.empty:
            def get_status(date_val):
                if date_val < today: return "🔴 Overdue"
                elif date_val <= today + datetime.timedelta(days=7): return "🟡 Upcoming"
                else: return "🟢 Scheduled"
            
            display_df = pm_df.copy()
            display_df['Status'] = display_df['next_due_date_dt'].apply(get_status)
            display_df = display_df.drop(columns=['next_due_date_dt'])
            st.dataframe(display_df, use_container_width=True)
            
            action_col1, action_col2 = st.columns(2)
            with action_col1:
                with st.form("generate_pm_wo_form"):
                    st.markdown("<div class='kpi-title'>Generate Work Order from Schedule</div>", unsafe_allow_html=True)
                    sched_to_exec = st.selectbox("Select Schedule:", display_df['Schedule ID'].tolist())
                    if st.form_submit_button("Generate Preventive Work Order"):
                        sched_row = display_df[display_df['Schedule ID'] == sched_to_exec].iloc[0]
                        existing_wo = get_all_work_orders()
                        new_wo_id = get_next_id(existing_wo, 'Work Order ID', 'WO')
                        
                        add_work_order(new_wo_id, sched_row['Product ID'], "Preventive Maintenance", "Medium", sched_row['Technician'], str(today), "Open", f"System Auto-Generated Preventive Maintenance.\nChecklist Items:\n{sched_row['Checklist']}")
                        
                        curr_due = pd.to_datetime(sched_row['Next Due Date']).date()
                        fq = sched_row['Frequency']
                        next_date = curr_due + datetime.timedelta(days=1 if fq=="Daily" else 7 if fq=="Weekly" else 30 if fq=="Monthly" else 90 if fq=="Quarterly" else 365 if fq=="Yearly" else 30)
                        
                        update_pm_schedule_date(sched_to_exec, str(next_date))
                        st.success(f"✅ Process Completed: Work Order `{new_wo_id}` generated from Schedule `{sched_to_exec}`. Next due date advanced to **{next_date}**.")
                        st.rerun()
            with action_col2:
                with st.form("delete_pm_form"):
                    st.markdown("<div class='kpi-title'>Delete Schedule</div>", unsafe_allow_html=True)
                    sched_to_del = st.selectbox("Select Schedule:", display_df['Schedule ID'].tolist())
                    confirm_pm_delete = st.checkbox("I understand this permanently deletes this schedule.")
                    if st.form_submit_button("Delete Schedule"):
                        if confirm_pm_delete:
                            delete_pm_schedule(sched_to_del)
                            st.warning(f"🗑️ Process Completed: Schedule `{sched_to_del}` has been permanently deleted.")
                            st.rerun()
                        else:
                            st.error("Please check the confirmation box before deleting.")
        else:
            st.info("No preventive maintenance schedules found. Create one in the next tab.")
            
        st.markdown("---")
        st.markdown("<h4 style='color:#FFFFFF;'>📜 Preventive Maintenance History</h4>", unsafe_allow_html=True)
        if not wo_df.empty: st.dataframe(wo_df[wo_df['Maintenance Type'] == 'Preventive Maintenance'][['Work Order ID', 'Product ID', 'Date Assigned', 'Status', 'Assigned Technician']], use_container_width=True)
        else: st.write("No historical records found.")
            
    with tab2:
        st.markdown("<h4 style='color:#FFFFFF;'>➕ Configure Maintenance Frequency & Schedule</h4>", unsafe_allow_html=True)
        machine_options = list(df['Product ID'].unique())
        freq_options = ["Daily", "Weekly", "Monthly", "Quarterly", "Yearly"]
        
        # FIX: default machine now prioritizes the shared target_machine_id so
        # it matches whatever was last selected in AI Maintenance Assistant or
        # the AI Recommendations tab. Falls back to pm_auto_fill if that ID
        # somehow isn't valid for this dataset.
        default_mac = st.session_state.get('target_machine_id', st.session_state['pm_auto_fill']['machine_id'])
        if default_mac not in machine_options:
            default_mac = st.session_state['pm_auto_fill']['machine_id']
        default_freq = st.session_state['pm_auto_fill']['frequency'].capitalize()
        m_index = machine_options.index(default_mac) if default_mac in machine_options else 0
        f_index = freq_options.index(default_freq) if default_freq in freq_options else 2
        with st.form("create_pm_schedule"):
            c1, c2 = st.columns(2)
            with c1:
                target_machine = st.selectbox("Target Machine ID (Product ID):", machine_options, index=m_index)
                frequency = st.selectbox("Maintenance Frequency:", freq_options, index=f_index)
            with c2:
                next_due = st.date_input("First / Next Due Date:", min_value=today)
                technician = st.selectbox("Assign Technician:", ["David", "Alex", "Swaraj", "Aditya", "John", "Sarah"])
            
            checklist = st.text_area("Maintenance Checklist (Items to inspect/replace):", value=st.session_state['pm_auto_fill']['checklist'], height=150)
            
            if st.form_submit_button("Save Schedule"):
                new_id = get_next_id(pm_df, 'Schedule ID', 'PM')
                add_pm_schedule(new_id, target_machine, frequency, str(next_due), technician, checklist)
                # Keep the canonical target machine ID aligned with the machine
                # this schedule was just created for.
                sync_target_machine(target_machine)
                st.success(f"✅ Schedule {new_id} is saved successfully!")
                st.rerun()
                
    with tab3:
        st.markdown("<h4 style='color:#FFFFFF;'>🤖 AI-Based Maintenance Recommendations</h4>", unsafe_allow_html=True)
        
        # FIX: default machine here is now driven by the shared target_machine_id
        # too, so it matches the AI Maintenance Assistant / Create Schedule
        # selection. Selecting a different machine here also updates the
        # canonical ID so Create Schedule picks it up on the next rerun.
        pm_machine_options = list(df['Product ID'].unique())
        pm_target_default = st.session_state.get('target_machine_id', pm_machine_options[0])
        pm_target_index = pm_machine_options.index(pm_target_default) if pm_target_default in pm_machine_options else 0
        ai_machine = st.selectbox("Select Machine for AI Analysis:", pm_machine_options, index=pm_target_index, key="ai_pm_mac")
        sync_target_machine(ai_machine)
        
        model_name_pm = st.text_input("Local Ollama Model Name:", value="llama3.2:1b", key="pm_model")
        
        if st.button("Generate Recommendations"):
            m_row = df[df['Product ID'] == ai_machine].iloc[0]
            prompt = f"""You are an industrial AI. Analyze this machine:
Product ID: {ai_machine}
Type: {m_row['Type']}
Tool Wear: {m_row['Tool wear [min]']} min
Past Failure: {'Yes' if m_row['Machine failure']==1 else 'No'}
Suggest a preventive maintenance schedule. Return ONLY valid JSON exactly like this:
{{
  "recommended_frequency": "Monthly",
  "reasoning": "Tool wear is approaching limits.",
  "checklist": "1. Replace bit\\n2. Check coolant\\n3. Verify RPM stability"
}}"""
            with st.spinner("Analyzing telemetry via Ollama..."):
                try:
                    res = requests.post('http://localhost:11434/api/generate', json={'model': model_name_pm, 'prompt': prompt, 'stream': False}, timeout=60)
                    if res.status_code == 200:
                        try:
                            text = res.json()['response']
                            rec_json = extract_json_from_llm(text)
                            if rec_json:
                                st.success(f"✅ Process Completed: AI recommendation generated for **{ai_machine}**.")

                                recommended_frequency = str(rec_json.get('recommended_frequency', 'Monthly'))
                                reasoning = str(rec_json.get('reasoning', 'No reasoning provided.'))
                                raw_checklist = rec_json.get('checklist', '')

                                if isinstance(raw_checklist, list):
                                    checklist_items = [str(item).strip() for item in raw_checklist if str(item).strip()]
                                else:
                                    checklist_text = str(raw_checklist).replace('\\n', '\n')
                                    checklist_items = []
                                    for line in checklist_text.split('\n'):
                                        cleaned = re.sub(r'^\s*(\d+[\.\)]|[-*])\s*', '', line).strip()
                                        if cleaned:
                                            checklist_items.append(cleaned)

                                st.markdown(f"""
                                <div class='card-panel' style='margin-top:10px;'>
                                    <div class='kpi-title'>Recommended Frequency</div>
                                    <div class='kpi-val' style='margin-bottom:14px;'>{html.escape(recommended_frequency)}</div>
                                    <div class='kpi-title'>Reasoning</div>
                                    <div style='color:#FFFFFF; margin-bottom:14px;'>{html.escape(reasoning)}</div>
                                    <div class='kpi-title'>Checklist</div>
                                </div>
                                """, unsafe_allow_html=True)

                                for idx, item in enumerate(checklist_items, start=1):
                                    st.markdown(f"**{idx}.** {item}")

                                formatted_chk = "\n".join(checklist_items) if checklist_items else str(raw_checklist)
                                st.session_state['pm_temp_rec'] = {'machine_id': ai_machine, 'frequency': recommended_frequency, 'checklist': formatted_chk}
                            else:
                                st.warning("Ollama returned text that could not be parsed as strict JSON.")
                                st.text(text)
                        except Exception as parse_err:
                            st.error(f"Could not read Ollama response: {parse_err}")
                    else:
                        st.error(f"Ollama returned status code {res.status_code}.")
                except requests.exceptions.ConnectionError:
                    st.error("⚠️ Could not connect to Ollama at http://localhost:11434. Make sure Ollama is running.")
                except requests.exceptions.Timeout:
                    st.error("⏱️ Request to Ollama timed out after 60s.")
                except Exception as e:
                    st.error(f"Unexpected error: {e}")
                    
        if 'pm_temp_rec' in st.session_state:
            st.markdown("---")
            if st.button("📋 Apply Recommendations to 'Create Schedule' Form"):
                st.session_state['pm_auto_fill'] = st.session_state['pm_temp_rec']
                sync_target_machine(st.session_state['pm_temp_rec']['machine_id'])
                st.success("✅ Process Completed: Recommendations applied. Switch over to the '➕ Create Schedule' tab to save your new maintenance plan.")
